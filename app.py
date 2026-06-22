import io
import os
import re
import tempfile
import subprocess

import fitz
import streamlit as st
from docx import Document


st.set_page_config(page_title="PDF OCR → Markdown / Word", page_icon="📄")

st.title("📄 PDF → Markdown / Word")
st.write("Sube un PDF. La app intentará hacer OCR si el PDF no tiene texto digital.")


def clean_text(text: str) -> str:
    text = text.replace("\x00", "")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = re.sub(r" +\n", "\n", text)
    return text.strip()


def extract_text_from_pdf(pdf_bytes: bytes) -> str:
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    pages = []

    for page_number, page in enumerate(doc, start=1):
        text = page.get_text("text")
        text = clean_text(text)

        if text:
            pages.append(f"# Página {page_number}\n\n{text}")

    return "\n\n".join(pages).strip()


def run_ocrmypdf(pdf_bytes: bytes) -> bytes:
    with tempfile.TemporaryDirectory() as tmpdir:
        input_path = os.path.join(tmpdir, "input.pdf")
        output_path = os.path.join(tmpdir, "output_ocr.pdf")

        with open(input_path, "wb") as f:
            f.write(pdf_bytes)

        command = [
            "ocrmypdf",
            "--force-ocr",
            "--deskew",
            "--clean",
            "--optimize", "1",
            "--language", "spa+eng",
            input_path,
            output_path,
        ]

        subprocess.run(command, check=True)

        with open(output_path, "rb") as f:
            return f.read()


def pdf_to_markdown(pdf_bytes: bytes, use_ocr: bool) -> str:
    text = extract_text_from_pdf(pdf_bytes)

    if not use_ocr and len(text) > 500:
        return text

    try:
        ocr_pdf_bytes = run_ocrmypdf(pdf_bytes)
        ocr_text = extract_text_from_pdf(ocr_pdf_bytes)

        if len(ocr_text) > len(text):
            return ocr_text

        return text

    except Exception as e:
        st.warning(
            "No se ha podido ejecutar OCR. Se usará el texto extraído directamente del PDF."
        )
        st.caption(str(e))
        return text


def improve_markdown_structure(markdown_text: str) -> str:
    lines = markdown_text.splitlines()
    cleaned = []

    for line in lines:
        line = line.strip()

        if not line:
            cleaned.append("")
            continue

        if re.match(r"^Página \d+", line, re.IGNORECASE):
            cleaned.append(f"# {line}")
            continue

        if re.match(r"^\d+\.\s+[A-ZÁÉÍÓÚÑ]", line):
            cleaned.append(f"## {line}")
            continue

        if re.match(r"^[A-ZÁÉÍÓÚÑ][A-ZÁÉÍÓÚÑ\s,.;:/()0-9-]{8,}$", line):
            cleaned.append(f"## {line.title()}")
            continue

        if re.match(r"^[-•]\s+", line):
            cleaned.append(f"- {re.sub(r'^[-•]\s+', '', line)}")
            continue

        cleaned.append(line)

    result = "\n".join(cleaned)
    result = re.sub(r"\n{3,}", "\n\n", result)
    return result.strip()


def markdown_to_word(markdown_text: str) -> bytes:
    document = Document()

    for line in markdown_text.splitlines():
        line = line.strip()

        if not line:
            document.add_paragraph()
        elif line.startswith("# "):
            document.add_heading(line.replace("# ", "", 1), level=1)
        elif line.startswith("## "):
            document.add_heading(line.replace("## ", "", 1), level=2)
        elif line.startswith("### "):
            document.add_heading(line.replace("### ", "", 1), level=3)
        elif line.startswith("- "):
            document.add_paragraph(line.replace("- ", "", 1), style="List Bullet")
        elif re.match(r"^\d+\.\s+", line):
            document.add_paragraph(re.sub(r"^\d+\.\s+", "", line), style="List Number")
        else:
            document.add_paragraph(line)

    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


uploaded_file = st.file_uploader("Sube tu PDF", type=["pdf"])

use_ocr = st.checkbox(
    "Forzar OCR para PDFs escaneados",
    value=True,
)

if uploaded_file is not None:
    pdf_bytes = uploaded_file.read()
    base_name = uploaded_file.name.rsplit(".", 1)[0]

    with st.spinner("Procesando PDF..."):
        markdown_text = pdf_to_markdown(pdf_bytes, use_ocr=use_ocr)
        markdown_text = improve_markdown_structure(markdown_text)
        word_bytes = markdown_to_word(markdown_text)

    st.success("Documento procesado.")

    st.subheader("Vista previa")
    st.text_area("Contenido extraído", markdown_text[:8000], height=350)

    col1, col2 = st.columns(2)

    with col1:
        st.download_button(
            label="📥 Descargar Markdown (.md)",
            data=markdown_text,
            file_name=f"{base_name}.md",
            mime="text/markdown",
        )

    with col2:
        st.download_button(
            label="📥 Descargar Word (.docx)",
            data=word_bytes,
            file_name=f"{base_name}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
